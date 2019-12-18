import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
def style_set(output):
 obj_styles = output.styles
 obj_charstyle = obj_styles.add_style('Style_00', WD_STYLE_TYPE.CHARACTER)
 obj_font = obj_charstyle.font
 obj_font.size = Pt(9)
 obj_font.name = 'Arial'
 obj_charstyle = obj_styles.add_style('Style_01', WD_STYLE_TYPE.CHARACTER)
 obj_font = obj_charstyle.font
 obj_font.size = Pt(10)
 obj_font.name = 'Arial'
 obj_charstyle = obj_styles.add_style('Style_10', WD_STYLE_TYPE.CHARACTER)
 obj_font = obj_charstyle.font
 obj_font.size = Pt(10)
 obj_font.name = 'Cambria'
 obj_charstyle = obj_styles.add_style('Style_11', WD_STYLE_TYPE.CHARACTER)
 obj_font = obj_charstyle.font
 obj_font.size = Pt(11)
 obj_font.name = 'Arial'
 obj_charstyle = obj_styles.add_style('Style_0', WD_STYLE_TYPE.CHARACTER)
 obj_font = obj_charstyle.font
 obj_font.size = Pt(9)
 obj_font.name = 'Calibri'
 obj_charstyle = obj_styles.add_style('Style_1', WD_STYLE_TYPE.CHARACTER)
 obj_font = obj_charstyle.font
 obj_font.size = Pt(10)
 obj_font.name = 'Calibri'
 obj_charstyle = obj_styles.add_style('Default', WD_STYLE_TYPE.CHARACTER)
 obj_font = obj_charstyle.font
 obj_font.size = Pt(11)
Page 9 of 16
 obj_font.name = 'Calibri'
def str2bin(text):
 binary_val = bin(int.from_bytes(text.encode(), 'big'))
 return binary_val
def bin2str(binary_val):
 binary_val = "0b" + binary_val
 n = int(binary_val, 2)
 text = n.to_bytes((n.bit_length() + 7) // 8, 'big').decode()
 return text
def encryptbinaryintotext(word, para, data):
 if data == "0":
 para.add_run(word, style='Style_0')
 if data == "1":
 para.add_run(word, style='Style_1')
 if data == "00":
 para.add_run(word, style='Style_00')
 if data == "01":
 para.add_run(word, style='Style_01')
 if data == "10":
 para.add_run(word, style='Style_10')
 if data == "11":
 para.add_run(word, style='Style_11')
def encrypt(text):
 input = docx.Document("Nature.docx")
 binary_data = str2bin(text)
 output = docx.Document()
 style_set(output)
 x = output.add_paragraph()
 j = 2
 word = ""
 for p in input.paragraphs:
 for r in p.runs:
 for i in r.text:
 word = word + i
 if i == ' ' or i == '.':
 data = binary_data[j:j+2]
 j = j + 2
 if data:
 encryptbinaryintotext(word, x, data)
 word = ""
 else:
 x.add_run(word, style='Default')
Page 10 of 16
 word = ""
 output.save("output.docx")
def decrypt():
 input = docx.Document("output.docx")
 data = ""
for p in input.paragraphs:
 for r in p.runs:
 if r.style.font.name == 'Cambria':
 data = data + "10"
 elif r.style.font.name == 'Calibri':
 x = r.style.font.size.pt
 if x == 9:
 data = data + "0"
 if x == 10:
 data = data + "1"
 if x == 11:
 return bin2str(data)
 elif r.style.font.name == 'Arial':
 x = r.style.font.size.pt
 if x == 9:
 data = data + "00"
 if x == 10:
 data = data + "01"
 if x == 11:
 data = data + "11"
 return data
def main():
 while True:
 action = input("Do you want to encrypt or decrypt the file (d/e) or exit: ")
 if action == 'e':
 text = input("Enter the data to encrypt: ")
 encrypt(text)
 elif action == 'd':
 text = decrypt()
 print("Decrypted code is: " + text )
 elif action == 'exit':
 break
 else:
 print("Invalid input try again")
if __name__ == '__main__':
 main()